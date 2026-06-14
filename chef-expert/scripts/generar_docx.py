#!/usr/bin/env python3
"""
Genera un documento DOCX profesional a partir de datos JSON de una receta.

Uso:
  echo '{
    "titulo": "Título",
    "raciones": 4,
    "ingredientes": [["500 g", "Ingrediente"], ...],
    "pasos": [["1. Título paso", "Cuerpo"], ...],
    "ajustes": [["Nombre", "Descripción"], ...],
    "consejos": ["Consejo 1", ...],
    "etiquetas": ["etiqueta1"]
  }' | python3 scripts/generar_docx.py --titulo "Título" --output "receta.docx"
"""

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def crear_docx(data, output_path):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data.get('titulo', 'Receta'))
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    partes = []
    raciones = data.get('raciones')
    if raciones:
        partes.append(f'{raciones} raciones')

    etiquetas = data.get('etiquetas', [])
    if etiquetas:
        partes.append(' | '.join(etiquetas))

    run = sub.add_run(' | '.join(partes))
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Ingredients
    doc.add_heading('Ingredientes', level=1)

    ingredientes = data.get('ingredientes', [])
    table = doc.add_table(rows=len(ingredientes), cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (cant, ing) in enumerate(ingredientes):
        row = table.rows[i]
        row.cells[0].text = cant
        row.cells[1].text = ing
        for cell in row.cells:
            for p in cell.paragraphs:
                p.style.font.size = Pt(10)

    doc.add_paragraph()

    # Steps
    doc.add_heading('Elaboración', level=1)

    for heading, body in data.get('pasos', []):
        h = doc.add_heading(heading, level=2)
        p = doc.add_paragraph(body)
        p.style.font.size = Pt(11)

    # Adjustments
    ajustes = data.get('ajustes', [])
    if ajustes:
        doc.add_heading('Micro-ajustes según tu gusto', level=1)
        doc.add_paragraph('Cada casa tiene su punto ideal. Aquí tienes cómo ajustar:')
        for adj_title, adj_body in ajustes:
            h = doc.add_heading(adj_title, level=2)
            p = doc.add_paragraph(adj_body)
            p.style.font.size = Pt(11)

    # Tips
    consejos = data.get('consejos', [])
    if consejos:
        doc.add_heading('Consejos adicionales', level=1)
        for tip in consejos:
            doc.add_paragraph(tip, style='List Bullet')

    doc.save(output_path)
    print(f'Documento guardado en: {output_path}')


def main():
    parser = argparse.ArgumentParser(description='Genera DOCX de receta')
    parser.add_argument('--titulo', '-t', required=True, help='Título de la receta')
    parser.add_argument('--output', '-o', required=True, help='Ruta del archivo .docx de salida')
    args = parser.parse_args()

    raw = sys.stdin.read()
    data = json.loads(raw)
    data['titulo'] = args.titulo

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    crear_docx(data, str(output_path))


if __name__ == '__main__':
    main()
