from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- Title ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Patatas guisadas con carrillera')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('4 raciones | Receta adaptada para olla a presión')
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# --- Ingredients ---
doc.add_heading('Ingredientes', level=1)

ingredients = [
    ('480 g', 'Carrillera de cerdo (troceada en bocados de ~4 cm)'),
    ('600 g', 'Patatas (partidas, no cortadas)'),
    ('100 g', 'Asadillo (½ bote pequeño)'),
    ('1 unidad', 'Zanahoria mediana'),
    ('1 cdta', 'Cebolla en polvo'),
    ('½ vaso (100 ml)', 'Vino blanco'),
    ('½ cdta', 'Pimentón de la Vera (dulce o picante al gusto)'),
    ('1 hoja', 'Laurel'),
    ('1 pizca', 'Hebras de azafrán (o colorante)'),
    ('-', 'Aceite de oliva virgen extra'),
    ('-', 'Sal'),
    ('-', 'Agua (para cubrir)'),
]

table = doc.add_table(rows=len(ingredients), cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (cant, ing) in enumerate(ingredients):
    row = table.rows[i]
    row.cells[0].text = cant
    row.cells[1].text = ing
    for cell in row.cells:
        for p in cell.paragraphs:
            p.style.font.size = Pt(10)

doc.add_paragraph()

# --- Elaboration ---
doc.add_heading('Elaboración', level=1)

steps = [
    ('1. Dorar la carne',
     'Cubre el fondo de la olla a presión con aceite de oliva y calienta a fuego fuerte. '
     'Trocea la carrillera en bocados de ~4 cm y dórala por todos lados hasta que esté bien sellada. '
     'Retírala y resérvala.'),

    ('2. Rehogar la verdura',
     'Baja a fuego medio. En la misma olla, añade la zanahoria pelada y picada fina. '
     'Rehoga 5 minutos. Mientras, rehidrata la cebolla en polvo en un chorrito de agua caliente '
     '(2 cucharadas) durante 5 minutos. Incorpora la cebolla rehidratada y el asadillo. '
     'Cocina 5 minutos más removiendo de vez en cuando.'),

    ('3. Sofreír y desglasar',
     'Aparta la olla del fuego un momento. Añade el pimentón y remueve rápido unos segundos '
     'para que no se queme. Vuelve al fuego, vierte el vino blanco y añade el laurel. '
     'Sube a fuego medio-alto y deja evaporar el alcohol 2 minutos. Reincorpora la carne reservada.'),

    ('4. Primera cocción (solo la carne)',
     'Cubre con agua justo hasta cubrir la carne. Tapa la olla y sube el fuego al máximo '
     'hasta que alcance presión (empiece a silbar). Entonces baja el fuego a medio-bajo para '
     'mantener la presión constante y cuenta:\n\n'
     '• Olla con selector de presión → Posición 1 → 10 minutos\n'
     '• Olla sin selector (presión única alta) → 8 minutos\n\n'
     'Pasado el tiempo, retira del fuego y deja salir la presión. '
     'Puedes usar salida natural (esperar) o forzada (grifo/válvula según tu olla).'),

    ('5. Añadir las patatas',
     'Pela las patatas, lávalas y pártejas con el cuchillo (no las cortes limpiamente — '
     'al romperlas sueltan más almidón y espesan el caldo de forma natural). '
     'Añádelas a la olla junto con las hebras de azafrán y sal al gusto. '
     'Si falta líquido, añade agua caliente hasta casi cubrir las patatas.'),

    ('6. Segunda cocción (con patatas)',
     'Tapa la olla y vuelve a subir al máximo hasta alcanzar presión. Baja el fuego '
     'a medio-bajo y cuenta:\n\n'
     '• Posición 1 → 8-10 minutos\n'
     '• Sin selector (presión alta) → 6-8 minutos\n\n'
     'Al destapar, la patata debe estar tierna y el caldo ligeramente espeso. '
     'Si ves demasiado caldo líquido, destapa y deja reducir 2-3 minutos a fuego medio.'),

    ('7. Reposo y servicio',
     'Deja reposar tapado 5 minutos fuera del fuego. Sirve caliente en plato hondo. '
     'Acompaña con pan para mojar en la salsa.'),
]

for heading, body in steps:
    h = doc.add_heading(heading, level=2)
    p = doc.add_paragraph(body)
    p.style.font.size = Pt(11)

# --- Micro-ajustes según gusto ---
doc.add_heading('Micro-ajustes según tu gusto', level=1)

doc.add_paragraph('Cada casa tiene su punto ideal. Aquí tienes cómo ajustar tiempos y texturas:')

adjustments = [
    ('Patata más entera / menos hecha',
     'Reduce el tiempo de la 2ª cocción: 6-7 min (presión 1) o 5 min (presión alta). '
     'Trocea las patatas en cascos más grandes (~6 cm).'),
    ('Patata más deshecha / caldo cremoso',
     'Aumenta el tiempo de 2ª cocción: 12 min (presión 1) o 10 min (presión alta). '
     'Además, aplasta 3-4 trozos de patata contra la pared de la olla con una cuchara de madera '
     'antes de servir y remueve para integrar el almidón.'),
    ('Carnera más melosa (que se deshaga)',
     'Alarga la 1ª cocción: 15 min (posición 1) o 12 min (presión alta). '
     'Córtala en trozos más pequeños (~3 cm) para que se impregne mejor.'),
    ('Carne con más mordida (al dente)',
     'Acorta la 1ª cocción: 6-7 min (posición 1) o 5 min (presión alta). '
     'Córtala en trozos grandes (~5 cm).'),
    ('Caldo más espeso',
     'Usa 500 g de patatas + 1 patata extra pequeña (70-80 g) que aplastarás antes de servir. '
     'También puedes añadir 1 cucharadita de harina de maíz (maicena) disuelta en agua fría '
     'al final y cocer 2 min sin presión.'),
    ('Caldo más ligero / suelto',
     'Usa 500 g de patatas en lugar de 600 g. No aplastes ninguna patata. '
     'Reduce el tiempo de 2ª cocción al mínimo indicado.'),
    ('Más sabor a pimentón',
     'Aumenta a 1 cucharadita rasa. Agrega la mitad al sofreír y la otra mitad al final, '
     'fuera del fuego, para que el sabor esté más fresco.'),
    ('Sin vino',
     'Sustituye el vino por 100 ml de caldo de carne o agua con 1 cucharada de vinagre de vino.'),
    ('Picante',
     'Usa pimentón picante en lugar de dulce, o añade ½ guindilla cayena junto con el vino.'),
]

for adj_title, adj_body in adjustments:
    h = doc.add_heading(adj_title, level=2)
    p = doc.add_paragraph(adj_body)
    p.style.font.size = Pt(11)

# --- Tips ---
doc.add_heading('Consejos adicionales', level=1)

tips = [
    'Tiempo total real: aproximadamente 45-50 minutos desde que empiezas hasta que sirves.',
    'Punto ideal de la carne: debe estar melosa, que se deshaga ligeramente al presionarla con el tenedor pero sin desintegrarse.',
    'Congelación: este guiso aguanta perfectamente 2-3 meses en el congelador. Se conserva bien el sabor y la textura.',
    'Si usas olla normal (no a presión): dobla los tiempos aproximados. La 1ª cocción ~60 min y la 2ª ~30 min.',
    'La carrillera se puede sustituir por costillas, jarrete, morcillo o pollo. Ajusta tiempos según el tipo de carne.',
]

for tip in tips:
    doc.add_paragraph(tip, style='List Bullet')

# --- Save ---
out_path = '/root/opencode-projects/receta-patatas-carrillera.docx'
doc.save(out_path)
print(f'Documento guardado en: {out_path}')
